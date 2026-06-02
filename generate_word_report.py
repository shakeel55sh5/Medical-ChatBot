from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_screenshot_placeholder(doc, description):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[ INSERT SCREENSHOT HERE ]\nDescription: {description}\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add a border-like effect using a single cell table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = f"\n\n\n\n\n\n\nPlace Screenshot: {description}\n\n\n\n\n\n\n"
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def create_word_report():
    doc = Document()

    # --- Title Page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("\n\n\nPAKISTANI MEDICAL AI CHATBOT\nUSING MACHINE LEARNING")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Final Year Project Report")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(20)

    for _ in range(5): doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("Submitted By\n")
    info_run.font.size = Pt(14)
    
    student_info = info_p.add_run("Student Name: ____________________\nRegistration No: ____________________\nDepartment: Computer Science\nUniversity: ____________________\nSession: ____________________")
    student_info.bold = True
    student_info.font.size = Pt(14)

    doc.add_paragraph()
    
    supervisor_p = doc.add_paragraph()
    supervisor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    supervisor_run = supervisor_p.add_run("Submitted To\n")
    supervisor_run.font.size = Pt(14)
    sup_name = supervisor_p.add_run("Supervisor: ____________________")
    sup_name.bold = True
    sup_name.font.size = Pt(14)

    doc.add_page_break()

    # --- Abstract ---
    heading = doc.add_heading("ABSTRACT", level=1)
    doc.add_paragraph(
        "The Pakistani Medical AI Chatbot is an intelligent healthcare assistance system developed using Python and Machine Learning techniques. "
        "The chatbot analyzes user-provided symptoms and predicts possible diseases using a Random Forest Classification model. "
        "The system then recommends suitable Pakistani medicine brands and suggests an appropriate doctor specialization.\n\n"
        "The primary objective of this project is to provide quick preliminary healthcare guidance to users while emphasizing that professional medical consultation remains essential. "
        "The system employs TF-IDF vectorization for text processing and Random Forest algorithms for disease classification.\n\n"
        "Keywords: Artificial Intelligence, Machine Learning, Healthcare Chatbot, Random Forest, TF-IDF, Disease Prediction, Python."
    )
    doc.add_page_break()

    # --- Acknowledgement ---
    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    doc.add_paragraph(
        "I would like to express my sincere gratitude to my supervisor for providing guidance, support, and valuable suggestions throughout the completion of this project. "
        "I am also thankful to the faculty members of the Computer Science Department for their encouragement and technical assistance.\n\n"
        "Special thanks to my family and friends for their continuous motivation and support."
    )
    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading("TABLE OF CONTENTS", level=1)
    sections = [
        "1. Introduction", "2. Problem Statement", "3. Objectives", "4. Literature Review",
        "5. System Analysis", "6. System Design", "7. Methodology", "8. Technologies Used",
        "9. Implementation", "10. Screenshots & Visual Evidence", "11. Testing and Results", "12. Advantages",
        "13. Limitations", "14. Future Enhancements", "15. Conclusion", "16. References",
        "Appendix A: Source Code", "Appendix B: Dataset", "Appendix C: User Manual"
    ]
    for section in sections:
        doc.add_paragraph(section)
    doc.add_page_break()

    # --- Main Content Sections ---
    content_data = [
        ("1. INTRODUCTION", "1.1 Background\nArtificial Intelligence has transformed many sectors including healthcare. Medical chatbots assist users in identifying potential health conditions based on symptoms.\n\n1.2 Purpose\nTo provide quick health guidance.\n\n1.3 Scope\nAccept symptoms, predict diseases, recommend medicines/doctors."),
        ("2. PROBLEM STATEMENT", "Healthcare facilities are not always immediately accessible. Many individuals seek preliminary information before consulting doctors. 'How can Machine Learning provide basic medical guidance?'"),
        ("3. OBJECTIVES", "Primary: Develop AI chatbot, predict diseases, recommend medicines.\nSecondary: Improve awareness, demonstrate ML applications."),
        ("4. LITERATURE REVIEW", "Review of existing medical AI and the choice of Random Forest for high accuracy and robustness in classification tasks."),
        ("5. SYSTEM ANALYSIS", "Existing systems are time-consuming and costly. The proposed AI chatbot provides instant, automated, and free preliminary guidance."),
        ("6. SYSTEM DESIGN", "Modular design including User Interface, TF-IDF Vectorizer, and Random Forest Intelligence engine."),
        ("7. METHODOLOGY", "Data collection, text preprocessing, feature extraction using TF-IDF, and model training with Random Forest."),
        ("8. TECHNOLOGIES USED", "Python, Scikit-Learn, python-docx, VS Code."),
        ("9. IMPLEMENTATION", "Implementation of the TfidfVectorizer and RandomForestClassifier for mapping symptoms to disease labels.")
    ]

    for title, text in content_data:
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)
        doc.add_page_break()

    # --- Screenshots Section ---
    doc.add_heading("10. SCREENSHOTS & VISUAL EVIDENCE", level=1)
    screenshots = [
        "Project Folder Structure",
        "Chatbot Welcome Screen",
        "User Symptom Input Example",
        "Disease Prediction Results",
        "Medicine & Doctor Recommendations",
        "Negation Detection Logic Test",
        "System Exit Screen"
    ]
    for sc in screenshots:
        add_screenshot_placeholder(doc, sc)
    doc.add_page_break()

    # --- Testing & Conclusion ---
    final_sections = [
        ("11. TESTING AND RESULTS", "Successful testing of disease prediction accuracy and negation handling ('no fever')."),
        ("12. ADVANTAGES", "Speed, 24/7 availability, localized Pakistani medicine recommendations."),
        ("13. LIMITATIONS", "Not a substitute for professional medical advice, limited disease database."),
        ("14. FUTURE ENHANCEMENTS", "Urdu language support, Mobile App, Real-time Hospital API."),
        ("15. CONCLUSION", "The project successfully demonstrates ML applications in the Pakistani healthcare context."),
        ("16. REFERENCES", "Scikit-learn docs, Machine Learning by Tom Mitchell, etc.")
    ]

    for title, text in final_sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)
        doc.add_page_break()

    # --- Appendices ---
    doc.add_heading("Appendix A: SOURCE CODE", level=1)
    with open("medical_bot.py", "r") as f:
        code = doc.add_paragraph().add_run(f.read())
        code.font.name = 'Courier New'
        code.font.size = Pt(8)
    doc.add_page_break()

    doc.add_heading("Appendix B: DATASET", level=1)
    with open("medical_data.py", "r") as f:
        data_code = doc.add_paragraph().add_run(f.read())
        data_code.font.name = 'Courier New'
        data_code.font.size = Pt(8)
    doc.add_page_break()

    doc.add_heading("Appendix C: USER MANUAL", level=1)
    doc.add_paragraph("1. Install Python.\n2. Install libs: pip install scikit-learn python-docx\n3. Run: python medical_bot.py\n4. Type 'exit' to quit.")

    output_path = "Medical_Chatbot_Report.docx"
    doc.save(output_path)
    print(f"Word document generated successfully: {output_path}")

if __name__ == "__main__":
    create_word_report()
